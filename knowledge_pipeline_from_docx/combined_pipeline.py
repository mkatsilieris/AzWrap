import os
import sys
import json
import re
import hashlib
import logging
import argparse
from io import BytesIO
from typing import List, Dict, Any
from tqdm import tqdm
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from dotenv import load_dotenv
from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import AzureError
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
import azure.search.documents.indexes.models as azsdim
from azure.search.documents.indexes.models import (
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    SimpleField,
    VectorSearch,
    VectorSearchProfile,
    ExhaustiveKnnAlgorithmConfiguration,
    ExhaustiveKnnParameters,
    SemanticConfiguration,
    SemanticPrioritizedFields,
    SemanticField,
    SemanticSearch,
)
from openai import AzureOpenAI

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# === Constants ===
# Files for tracking processing progress
CHECKPOINT_FILE = "processed_files.txt"  # Tracks successfully processed files
FAILED_FILE_LOG = "failed_files.txt"     # Tracks files that failed processing

# Global variable for clients, initialized later
clients = None

#############################################################
# CONFIGURATION LOADING FUNCTION
#############################################################

def load_configuration():
    """
    Loads configuration from environment variables using dotenv.
    Validates required variables and returns them as a dictionary.
    """
    print("🔧 Loading configuration from .env file...")
    load_dotenv()

    config = {
        "AZURE_TENANT_ID": os.getenv("AZURE_TENANT_ID"),
        "AZURE_CLIENT_ID": os.getenv("AZURE_CLIENT_ID"),
        "AZURE_CLIENT_SECRET": os.getenv("AZURE_CLIENT_SECRET"),
        "AZURE_SUBSCRIPTION_ID": os.getenv("AZURE_SUBSCRIPTION_ID"),
        "RESOURCE_GROUP": os.getenv("RESOURCE_GROUP"),
        "TARGET_ACCOUNT_NAME": os.getenv("TARGET_ACCOUNT_NAME"), # OpenAI Account
        "ACCOUNT_NAME": os.getenv("ACCOUNT_NAME"), # Storage Account
        "CONTAINER_NAME": os.getenv("CONTAINER_NAME"),
        "SEARCH_SERVICE_NAME": os.getenv("SEARCH_SERVICE_NAME"),
        "CORE_INDEX_NAME": os.getenv("CORE_INDEX_NAME"),
        "DETAILED_INDEX_NAME": os.getenv("DETAILED_INDEX_NAME"),
        "MODEL_NAME": os.getenv("MODEL_NAME"), # OpenAI Model
        # Consider adding TEMP_PATH and FORMAT_JSON_PATH here if they need to be configurable
        "TEMP_PATH": r"C:\Users\agkithko\OneDrive - Netcompany\Desktop\AzWrap-1\temp_json", # Default, consider making configurable
        "FORMAT_JSON_PATH": r"C:\Users\agkithko\OneDrive - Netcompany\Desktop\AzWrap-1\knowledge_pipeline_from_docx\format.json" # Default, consider making configurable
    }

    # Validate essential variables
    required_vars = [
        "AZURE_TENANT_ID", "AZURE_CLIENT_ID", "AZURE_CLIENT_SECRET", "AZURE_SUBSCRIPTION_ID",
        "RESOURCE_GROUP", "ACCOUNT_NAME", "CONTAINER_NAME", "SEARCH_SERVICE_NAME",
        "CORE_INDEX_NAME", "DETAILED_INDEX_NAME", "TARGET_ACCOUNT_NAME", "MODEL_NAME",
        "TEMP_PATH", "FORMAT_JSON_PATH"
    ]
    missing_vars = [var for var in required_vars if not config.get(var)]
    if missing_vars:
        logger.error(f"❌ Error: Missing required environment variables or config: {', '.join(missing_vars)}")
        sys.exit(1) # Exit if configuration is incomplete

    logger.info("✅ Configuration loaded successfully.")
    return config

#############################################################
# UTILITY FUNCTIONS (from main.py)
#############################################################

def load_format(format_path):
    """
    Loads the JSON format file that defines how documents should be processed.

    Args:
        format_path (str): Path to the JSON format file

    Returns:
        dict: The loaded JSON format configuration
    """
    print("📂 Loading JSON format file...")
    try:
        with open(format_path, "r", encoding="utf-8") as file:
            return json.load(file)
    except FileNotFoundError:
        logger.error(f"❌ Error: Format JSON file not found at {format_path}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        logger.error(f"❌ Error decoding format JSON file: {e}")
        sys.exit(1)

def delete_files_in_directory(directory_path):
    """
    Deletes all files in a given directory to clean up temporary files.

    Args:
        directory_path (str): Path to the directory to clean
    """
    print(f"🗑️ Starting file deletion in directory: {directory_path}")
    if os.path.exists(directory_path) and os.path.isdir(directory_path):
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                    logger.info(f"   ✔️ Deleted: {file_path}")
                # Optionally handle directories if needed
            except Exception as e:
                logger.error(f'   ❌ Failed to delete {file_path}. Reason: {e}')
    else:
        logger.warning(f"⚠️ Directory '{directory_path}' does not exist or is not a directory.")
    print("🏁 File deletion process completed")

def get_all_files_in_directory(directory_path):
    """
    Returns a list of all file paths in a given directory, excluding 'Metadata' files.

    Args:
        directory_path (str): Directory to scan for files

    Returns:
        list: List of file paths found in the directory
    """
    print(f"📂 Scanning directory for files: {directory_path}")
    file_list = []
    if not os.path.isdir(directory_path):
        logger.warning(f"⚠️ Directory '{directory_path}' does not exist.")
        return file_list

    for root, _, files in os.walk(directory_path):
        for file in files:
            if "Metadata" not in file: # Simple check
                file_path = os.path.join(root, file)
                file_list.append(file_path)
                logger.debug(f"   📄 Found file: {file_path}") # Use debug for less verbose output
            else:
                logger.debug(f"   🚫 Skipping metadata file: {file}")
    print(f"🔢 Total files found (excluding metadata): {len(file_list)}")
    return file_list

#############################################################
# CHECKPOINTING FUNCTIONS (from main.py)
#############################################################

def load_processed_files():
    """
    Load the list of previously processed files from checkpoint file.

    Returns:
        set: Set of processed file names
    """
    if not os.path.exists(CHECKPOINT_FILE):
        return set()
    try:
        with open(CHECKPOINT_FILE, "r", encoding="utf-8") as f:
            return set(line.strip() for line in f if line.strip())
    except Exception as e:
        logger.warning(f"⚠️ Error loading processed files checkpoint {CHECKPOINT_FILE}: {e}")
        return set()

def save_processed_file(filename):
    """
    Save a successfully processed file to the checkpoint file.

    Args:
        filename (str): Name of the processed file
    """
    try:
        with open(CHECKPOINT_FILE, "a", encoding="utf-8") as f:
            f.write(filename + "\n")
    except Exception as e:
        logger.warning(f"⚠️ Error saving processed file checkpoint {filename}: {e}")

def log_failed_file(filename_with_error):
    """
    Log a failed file for retry later.

    Args:
        filename_with_error (str): Name of the failed file along with error info
    """
    try:
        with open(FAILED_FILE_LOG, "a", encoding="utf-8") as f:
            f.write(filename_with_error + "\n")
    except Exception as e:
        logger.warning(f"⚠️ Error logging failed file {filename_with_error}: {e}")

def load_failed_files():
    """
    Load the list of previously failed files.

    Returns:
        set: Set of failed file names (or names with error context)
    """
    if not os.path.exists(FAILED_FILE_LOG):
        return set()
    try:
        with open(FAILED_FILE_LOG, "r", encoding="utf-8") as f:
            return set(line.strip() for line in f if line.strip())
    except Exception as e:
        logger.warning(f"⚠️ Error loading failed files log {FAILED_FILE_LOG}: {e}")
        return set()

#############################################################
# AZURE CLIENT INITIALIZATION (Adapted from main.py)
#############################################################
# Assuming AzWrap is installed or available in the path
# Add basic error handling for the AzWrap import
try:
    # === Dynamic Path Setup ===
    # Ensure AzWrap is importable. Adjust if AzWrap is installed as a package.
    CURRENT_DIR_FOR_AZWRAP = os.path.dirname(os.path.abspath(__file__))
    AZWRAP_ROOT = os.path.abspath(os.path.join(CURRENT_DIR_FOR_AZWRAP, ".."))
    if AZWRAP_ROOT not in sys.path:
        sys.path.append(AZWRAP_ROOT)

    from AzWrap.wrapper import (
        Identity, Subscription, ResourceGroup,
        SearchService, SearchIndex as AzWrapSearchIndex, AIService, OpenAIClient,
        StorageAccount # Assuming StorageAccount needed for blob operations directly via AzWrap
    )
except ImportError as e:
    logger.critical(f"CRITICAL ERROR: AzWrap modules could not be imported: {e}")
    logger.critical("Please ensure the AzWrap library is correctly installed or the path is set.")
    sys.exit(1)

def initialize_clients(config: Dict[str, Any]):
    """
    Initializes all necessary Azure clients (Identity, Subscription, Storage, Search, OpenAI).

    Reads configuration from the provided dictionary. Exits if essential configuration is missing.

    Args:
        config (dict): Dictionary containing configuration values loaded from .env.

    Returns:
        dict: A dictionary containing initialized client objects, or None on failure.
              Keys: "identity", "sub", "rg", "storage_account", "container_client",
                    "search_service", "search_index_client", "search_client_core",
                    "search_client_detail", "azure_oai_client"
    """
    print("🔧 Initializing Azure Clients...")
    clients_dict = {}
    try:
        # 1. Identity and Subscription
        print("   Authenticating and getting subscription...")
        clients_dict["identity"] = Identity(tenant_id=config["AZURE_TENANT_ID"], client_id=config["AZURE_CLIENT_ID"], client_secret=config["AZURE_CLIENT_SECRET"])
        subscription_info = clients_dict["identity"].get_subscription(config["AZURE_SUBSCRIPTION_ID"])
        clients_dict["sub"] = Subscription(identity=clients_dict["identity"], subscription=subscription_info, subscription_id=subscription_info.subscription_id)
        logger.info(f"      ✔️ Subscription '{clients_dict['sub'].subscription.display_name}' obtained.")

        # 2. Resource Group
        print(f"   Getting resource group '{config['RESOURCE_GROUP']}'...")
        clients_dict["rg"] = clients_dict["sub"].get_resource_group(config["RESOURCE_GROUP"])
        logger.info(f"      ✔️ Resource group obtained.")

        # 3. Storage Account and Container Client (using AzWrap if available, else standard SDK)
        print(f"   Getting storage account '{config['ACCOUNT_NAME']}' and container '{config['CONTAINER_NAME']}'...")
        # Using AzWrap StorageAccount
        clients_dict["storage_account"] = clients_dict["rg"].get_storage_account(config['ACCOUNT_NAME'])
        # Directly get the blob container client needed for downloads
        clients_dict["container_client"] = clients_dict["storage_account"].get_container_client(config['CONTAINER_NAME'])
        logger.info(f"      ✔️ Storage container client obtained.")


        # 4. Search Service and Clients
        print(f"   Getting search service '{config['SEARCH_SERVICE_NAME']}' and clients...")
        clients_dict["search_service"] = clients_dict["sub"].get_search_service(config["SEARCH_SERVICE_NAME"])
        # Get client for managing indexes
        clients_dict["search_index_client"] = clients_dict["search_service"].get_index_client()
        # Get clients for interacting with specific indexes
        # Use AzWrapSearchIndex (renamed from AzWrap) to get SearchClient instances
        core_index_azwrap: AzWrapSearchIndex = clients_dict["search_service"].get_index(config["CORE_INDEX_NAME"])
        detail_index_azwrap: AzWrapSearchIndex = clients_dict["search_service"].get_index(config["DETAILED_INDEX_NAME"])
        clients_dict["search_client_core"] = core_index_azwrap.get_search_client()
        clients_dict["search_client_detail"] = detail_index_azwrap.get_search_client()
        logger.info(f"      ✔️ Search index client and clients for '{config['CORE_INDEX_NAME']}', '{config['DETAILED_INDEX_NAME']}' obtained.")

        # 5. Azure OpenAI Client
        print(f"   Getting Azure OpenAI client for account '{config['TARGET_ACCOUNT_NAME']}'...")
        cog_mgmt_client = clients_dict["sub"].get_cognitive_client()
        account_details = next(
            (acc for acc in cog_mgmt_client.accounts.list_by_resource_group(clients_dict["rg"].azure_resource_group.name)
             if acc.name == config['TARGET_ACCOUNT_NAME']), None
        )
        if not account_details:
             raise ValueError(f"Azure OpenAI account '{config['TARGET_ACCOUNT_NAME']}' not found in resource group '{config['RESOURCE_GROUP']}'.")

        ai_service = AIService(clients_dict["rg"], cog_mgmt_client, account_details)
        clients_dict["azure_oai_client"] = ai_service.get_AzureOpenAIClient(api_version='2024-05-01-preview') # Specific API version
        logger.info(f"      ✔️ Azure OpenAI client obtained.")

        logger.info("✅ All Azure clients initialized successfully.")
        return clients_dict

    except Exception as e:
        logger.error(f"❌ Failed to initialize Azure clients: {e}")
        logger.exception("Client initialization failed.") # Log stack trace
        return None # Indicate failure

#############################################################
# DOC PARSING CLASS (from doc_parsing.py)
#############################################################

class DocParsing:
    def __init__(self, doc_instance, client: AzureOpenAI, json_format: dict, domain: str, sub_domain: str, model_name: str, doc_name: str):
        """
        Initialize the DocParsing class.

        Parameters:
            doc_instance: python-docx Document object to be parsed
            client: Azure OpenAI client for AI processing
            json_format: Template for the JSON structure
            domain: Domain category for the document
            sub_domain: Sub-domain category for the document
            model_name: Name of the AI model to use
            doc_name: Name of the document being processed (without extension)
        """
        print(f"🚀 Initializing DocParsing for document: {doc_name}")
        self.client = client
        self.doc = doc_instance
        self.format = json_format # Use the passed dictionary directly
        self.domain = domain
        self.model_name = model_name
        self.sub_domain = sub_domain
        self.doc_name = doc_name # Store the name without extension

    def _get_section_header_lines(self, section):
        """Helper to extract text lines from a section's header."""
        try:
            if not section or not section.header:
                return []

            lines = []
            # Gather paragraph text from the header
            for paragraph in section.header.paragraphs:
                txt = paragraph.text.strip()
                if txt:
                    lines.append(txt)

            # Gather table cell text from the header (if any)
            for table in section.header.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells)) # Join cell text for table lines
            return lines
        except Exception as e:
            logger.warning(f"   ⚠️ Error extracting header lines for section: {e}")
            return []

    def _parse_header_lines(self, header_lines):
        """Helper to parse header lines to extract the process title."""
        if not header_lines:
            return "Metadata" # Default if no lines or only empty lines

        # Pattern for process numbers (e.g., 1., 1.1., 1.1.1.)
        number_pattern = re.compile(r'^\d+(\.\d+)*\.$')
        # Pattern for specific metadata lines to ignore (Example from doc_parsing.py)
        meta_patterns = [r'^Εκδ\.', r'Σελ\.'] # Add more patterns if needed

        potential_title = "Metadata" # Start with default

        for i, line in enumerate(header_lines):
            line_stripped = line.strip()
            if not line_stripped: continue # Skip empty lines

            # Skip known metadata lines
            if any(re.search(pattern, line_stripped) for pattern in meta_patterns):
                continue

            # Check if line matches "Number.\tTitle" format
            if "\t" in line_stripped:
                parts = line_stripped.split("\t", 1)
                potential_num = parts[0].strip()
                potential_title_part = parts[1].strip() if len(parts) > 1 else ""
                if number_pattern.match(potential_num) and potential_title_part:
                    return potential_title_part # Found title directly

            # Check if line is just a process number
            elif number_pattern.match(line_stripped):
                # Look for a non-metadata title in the *next* non-empty line
                if i + 1 < len(header_lines):
                    next_line_stripped = header_lines[i+1].strip()
                    if next_line_stripped and not any(re.search(pattern, next_line_stripped) for pattern in meta_patterns):
                         # Check if the next line looks like a title (heuristic: doesn't start with a number pattern)
                        if not number_pattern.match(next_line_stripped.split()[0] if next_line_stripped else ""):
                            return next_line_stripped # Found title on the next line

            # If the line is not metadata and not a number, consider it a potential title
            # This handles cases where title appears alone without a preceding number line
            elif potential_title == "Metadata": # Only take the first potential title
                 potential_title = line_stripped


        # If no specific pattern matched, return the first non-metadata line found, or "Metadata"
        return potential_title

    def _extract_header_info(self, section):
        """Extracts process title from a section header."""
        try:
            lines = self._get_section_header_lines(section)
            header_title = self._parse_header_lines(lines)
            return header_title
        except Exception as e:
            logger.warning(f"   ⚠️ Error extracting header info: {e}")
            return "Unknown Header" # Return a default on error

    def _iterate_block_items_with_section(self, doc):
        """Iterates through document blocks (paragraphs, tables) yielding (section_index, block)."""
        # Logic adapted from combined_pipeline.py, seems more robust than original doc_parsing.py
        parent_elm = doc._element.body
        current_section_index = 0
        last_element_was_sectPr = False

        for child in parent_elm.iterchildren():
            if child.tag.endswith("p"):
                paragraph = Paragraph(child, doc)
                is_section_end_paragraph = bool(child.xpath("./w:pPr/w:sectPr"))
                if not is_section_end_paragraph:
                     yield current_section_index, paragraph
                if is_section_end_paragraph:
                    current_section_index += 1
                    last_element_was_sectPr = True
                else:
                    last_element_was_sectPr = False
            elif child.tag.endswith("tbl"):
                table = Table(child, doc)
                yield current_section_index, table
                last_element_was_sectPr = False
            elif child.tag.endswith('sectPr') and not last_element_was_sectPr:
                 current_section_index += 1

    def _extract_table_data(self, table):
        """Extracts text data from a table, joining cells with ' - '."""
        data = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                data.append(' - '.join(row_cells))
        return '\n'.join(data) # Join rows with newline

    def _is_single_process(self):
        """Checks if the document contains a single process based on headers."""
        print("   🔎 Checking document for single vs. multi-process structure...")
        section_headers = set()
        first_meaningful_header = None

        if not self.doc.sections:
            logger.warning("   ⚠️ Document has no sections.")
            return True, self.doc_name # Treat as single process with doc name as title

        for section_index, section in enumerate(self.doc.sections):
            header_title = self._extract_header_info(section)
            if header_title and header_title != "Metadata" and header_title != "Unknown Header":
                section_headers.add(header_title)
                if first_meaningful_header is None:
                    first_meaningful_header = header_title # Store the first valid header found

        num_unique_headers = len(section_headers)
        logger.info(f"   ℹ️ Found {num_unique_headers} unique meaningful header(s): {section_headers if section_headers else 'None'}")

        if num_unique_headers <= 1: # Treat 0 or 1 unique headers as single process
            title = first_meaningful_header if first_meaningful_header else self.doc_name
            logger.info(f"   ➡️ Document treated as single process with title: '{title}'")
            return True, title
        else:
            logger.info("   ➡️ Document identified as multi-process.")
            return False, None # Title is None for multi-process

    def extract_data(self):
        """
        Extracts content from the document, handling single/multi-process structures.

        Returns:
            dict: Keys are formatted headers/process names, values are extracted content strings.
        """
        print("   ⛏️ Extracting data based on document structure...")
        data_dict = {}
        is_single, process_title = self._is_single_process()

        if is_single:
            safe_title = re.sub(r'[\\/*?:"<>|]', '_', process_title) # Basic sanitization
            header_key = f"{safe_title}_single_process"
            print(f"   🏗️ Building content for single process: '{header_key}'")
            data_dict[header_key] = []
            for _, block in self._iterate_block_items_with_section(self.doc):
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    if text: data_dict[header_key].append(text)
                elif isinstance(block, Table):
                    table_text = self._extract_table_data(block)
                    if table_text: data_dict[header_key].append(table_text)
        else:
            print("   🏗️ Building content for multi-process document...")
            last_section_index = -1
            current_header_key = None
            current_section_content = []

            for section_index, block in self._iterate_block_items_with_section(self.doc):
                if section_index > last_section_index:
                    if current_header_key and current_section_content:
                         data_dict[current_header_key] = "\n".join(current_section_content) # Join previous section
                         logger.debug(f"      ✔️ Finalized content for section {last_section_index}: '{current_header_key}' ({len(current_section_content)} blocks)")

                    if section_index < len(self.doc.sections):
                         header_title = self._extract_header_info(self.doc.sections[section_index])
                         if not header_title or header_title == "Metadata":
                             header_title = f"Unknown_Section_{section_index}"
                         safe_header = re.sub(r'[\\/*?:"<>|]', '_', header_title)
                         current_header_key = f"{self.doc_name}_header_{safe_header}"
                         logger.debug(f"   New Section {section_index}: Header='{header_title}', Key='{current_header_key}'")
                         if current_header_key not in data_dict:
                              data_dict[current_header_key] = []
                         current_section_content = [] # Reset buffer
                    else:
                         logger.warning(f"   ⚠️ Warning: Block referenced section_index {section_index} > section count {len(self.doc.sections)}. Using last header '{current_header_key}'.")

                    last_section_index = section_index

                block_text = ""
                if isinstance(block, Paragraph):
                    block_text = block.text.strip()
                elif isinstance(block, Table):
                    block_text = self._extract_table_data(block)

                if block_text and current_header_key:
                     # Append text directly to the list in the dictionary
                     if current_header_key in data_dict:
                         data_dict[current_header_key].append(block_text)
                     else:
                         # This case might happen if the first block has no preceding header info
                         logger.warning(f"   ⚠️ Warning: No current header key for block, text ignored: '{block_text[:50]}...'")


            # Finalize the very last section after the loop
            if current_header_key and current_section_content:
                data_dict[current_header_key] = "\n".join(current_section_content)
                logger.debug(f"      ✔️ Finalized content for last section {last_section_index}: '{current_header_key}' ({len(current_section_content)} blocks)")


        # Join the collected content lines for each key into final strings
        final_data = {key: "\n".join(content_list).strip() for key, content_list in data_dict.items()}
        print(f"   ✅ Data extraction complete. Found {len(final_data)} process/section block(s).")
        return final_data

    def update_json_with_ai(self, content_to_parse: str, process_identifier: str):
        """
        Uses AI to parse document content into the structured JSON format.

        Parameters:
            content_to_parse: The text content extracted for a specific process/section.
            process_identifier: The identifier (like header key) for this process/section.

        Returns:
            str: JSON string containing the parsed content, or None on failure.
        """
        print(f"   🤖 Requesting AI to parse content for: '{process_identifier}' using model '{self.model_name}'...")
        format_str = json.dumps(self.format, indent=4, ensure_ascii=False)

        # Prompt emphasizing extraction, structure, and JSON-only output
        prompt = (
            "Parse the provided information about a specific process from the document and fill in the JSON structure below. "
            "Do not summarize, omit, or modify any details. Simply extract and organize the provided data into the corresponding fields of the JSON. "
            "There are more than one step and you have to include all of them.The step description has to be the whole text till the next step name"
            "Ensure every relevant detail is included without altering the content. "
            "The JSON format should follow this structure and include all fields, even if some of them are not present in the content (leave them empty or null if necessary):\n"
            f"{format_str}\n\n"
            "To make it clear the content you generate will be ONLY THE CONTENT of a json no \\n nothing.The first character {{ and the last character should be }}" # Escaped braces
            "Your response should be ONLY a JSON file content ready to be stored as json without other processing, with the exact format as shown above."
         )

        try:
            if not self.client:
                 logger.error("   ❌ Error: Azure OpenAI client is not initialized.")
                 return None

            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": f"Document Source Name: {self.doc_name}\nProcess/Section Identifier: {process_identifier}\n\nContent to Parse:\n---\n{content_to_parse}\n---"}
            ]

            output_llm = self.client.chat.completions.create(
                model=self.model_name,
                messages=messages,
                temperature=0,
                response_format={"type": "json_object"} # Request JSON output
            )

            ai_response_content = output_llm.choices[0].message.content
            logger.info(f"   ✔️ AI response received ({len(ai_response_content)} chars).")

            # Basic validation: Check if it looks like JSON
            if ai_response_content and ai_response_content.strip().startswith("{") and ai_response_content.strip().endswith("}"):
                return ai_response_content.strip()
            else:
                logger.warning(f"   ⚠️ Warning: AI response does not look like valid JSON: {ai_response_content[:100]}...")
                # Attempt to extract JSON if response_format failed or wasn't respected
                match = re.search(r'\{.*\}', ai_response_content, re.DOTALL)
                if match:
                    logger.info("   ℹ️ Extracted potential JSON from response.")
                    return match.group(0)
                else:
                    logger.error("   ❌ Error: Could not extract JSON object from AI response.")
                    return None

        except Exception as e:
            logger.error(f"   ❌ Error during AI call for '{process_identifier}': {e}")
            logger.exception("AI call failed.") # Log stack trace
            return None


    def _process_and_save_json(self, ai_json_string: str, output_path: str):
        """
        Processes the AI response string, validates JSON, adds metadata, and saves to file.

        Parameters:
            ai_json_string: Raw JSON string from the AI.
            output_path: Full path to save the output JSON file.
        """
        print(f"   💾 Processing AI response and saving to: {output_path}")
        try:
            # Parse the AI's JSON string into a Python dictionary
            json_data = json.loads(ai_json_string)

            # Create a new ordered dictionary to control the final structure
            ordered_data = {}
            ordered_data["doc_name"] = self.doc_name
            ordered_data["process_name"] = json_data.get("process_name", "Unknown Process Name")
            ordered_data["domain"] = self.domain
            ordered_data["subdomain"] = self.sub_domain

            # Add the rest of the fields from the AI response
            for key, value in json_data.items():
                if key not in ordered_data:
                    ordered_data[key] = value

            with open(output_path, 'w', encoding='utf-8') as file:
                json.dump(ordered_data, file, indent=4, ensure_ascii=False)

            logger.info(f"   ✅ JSON data successfully processed and written to {output_path}")

        except json.JSONDecodeError as e:
            logger.error(f"   ❌ Error decoding JSON from AI response: {e}")
            logger.error(f"      Raw AI response snippet: {ai_json_string[:200]}...")
        except Exception as e:
            logger.error(f"   ❌ Unexpected error processing/saving JSON for {output_path}: {e}")
            logger.exception("JSON saving failed.")


    def doc_to_json(self, output_dir="temp_json"):
        """
        Main method: Converts the document content into one or more structured JSON files.

        Parameters:
            output_dir: Directory to save the output JSON files.
        """
        print(f"🚢 Starting document-to-JSON conversion for '{self.doc_name}'...")

        extracted_data_dict = self.extract_data()

        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
                logger.info(f"   📁 Created output directory: {output_dir}")
            except OSError as e:
                logger.error(f"   ❌ Error creating output directory {output_dir}: {e}")
                return

        if not extracted_data_dict:
             logger.warning("   ⚠️ No data extracted from the document. Skipping JSON generation.")
             return

        for process_key, content in extracted_data_dict.items():
            print(f"\n   Processing section/process key: '{process_key}'")
            if "Metadata" in process_key or not content.strip():
                logger.info("      🚫 Skipping metadata section or empty content.")
                continue

            filename_base = process_key.split('_header_')[-1] if '_header_' in process_key else process_key.replace('_single_process', '')
            safe_filename_base = re.sub(r'[\\/*?:"<>| ]', '_', filename_base)
            max_len = 100 # Max filename length
            safe_filename_base = safe_filename_base[:max_len] if len(safe_filename_base) > max_len else safe_filename_base

            output_json_path = os.path.join(output_dir, f"{safe_filename_base}.json")
            print(f"      Target JSON file: {output_json_path}")

            ai_json_result = self.update_json_with_ai(content, process_key)

            if ai_json_result:
                self._process_and_save_json(ai_json_result, output_json_path)
            else:
                logger.error(f"      ⚠️ AI parsing failed for '{process_key}'. JSON file will not be created.")
                log_failed_file(f"{self.doc_name} - Section '{process_key}' - AI Parsing Failed")


        print(f"\n🏁 Document-to-JSON conversion completed for '{self.doc_name}'. Check '{output_dir}' for results.")
        return # Explicit return

#############################################################
# JSON PROCESSING CLASS (from json_processing.py)
#############################################################

class ProcessHandler:
    def __init__(self, json_path: str):
        """
        Initializes the class with the path to a JSON file containing process data.
        """
        print(f"🗂️ Initializing ProcessHandler for: {json_path}")
        if not os.path.exists(json_path):
            raise FileNotFoundError(f"JSON file not found at path: {json_path}")

        try:
            with open(json_path, 'r', encoding='utf-8') as file:
                self.json_data = json.load(file)
            logger.info(f"   ✔️ JSON file loaded successfully.")
            self.doc_name = self.json_data.get('doc_name', 'N/A')
            self.process_name = self.json_data.get('process_name', 'N/A')
            logger.info(f"   📄 Document Name: {self.doc_name}")
            logger.info(f"   📋 Process Name: {self.process_name}")

        except json.JSONDecodeError as e:
            logger.error(f"   ❌ Error decoding JSON from file: {e}")
            raise
        except Exception as e:
            logger.error(f"   ❌ An unexpected error occurred loading JSON: {e}")
            raise

    def _generate_hash_id(self, *args) -> str:
        """Generates a SHA-256 hash string from input arguments."""
        hasher = hashlib.sha256()
        for arg in args:
            hasher.update(str(arg).encode('utf-8'))
        return hasher.hexdigest()

    def generate_process_id(self) -> str:
        """
        Generates a unique string ID for the process based on its name and document name.
        """
        logger.debug(f"   🔢 Generating Process ID for '{self.process_name}' from doc '{self.doc_name}'")
        process_id = self._generate_hash_id(self.doc_name, self.process_name)
        logger.debug(f"      ✅ Generated Process ID: {process_id}")
        return process_id

    def generate_step_id(self, step_number: int, step_name: str, step_content: str) -> str:
        """
        Generates a unique string ID for a step based on process, step number, name, and content.
        """
        content_snippet = (step_content or "")[:100] # Use snippet for stability
        step_id = self._generate_hash_id(self.process_name, step_number, step_name, content_snippet)
        return step_id

    def _prepare_core_record(self, process_id: str) -> Dict[str, Any]:
        """
        Prepares the record for the core Azure Search index.
        """
        print("   📊 Preparing Core Record...")

        steps_info = [f"Step {s.get('step_number', 'N/A')}: {s.get('step_name', 'Unnamed')}" for s in self.json_data.get('steps', [])]
        summary_parts = [
            f"Process: {self.process_name}",
            f"Source Document: {self.doc_name}",
            f"Domain: {self.json_data.get('domain', 'N/A')}/{self.json_data.get('subdomain', 'N/A')}", # Changed from 'subdomain' to match doc_parsing
            f"Introduction: {self.json_data.get('introduction', '').strip()}",
            f"Short Description: {self.json_data.get('short_description', '').strip()}",
            f"Steps Overview:\n- " + "\n- ".join(steps_info) if steps_info else "No steps listed.",
            f"Related Products: {', '.join(self.json_data.get('related_products', [])) or 'None'}",
            f"Reference Documents: {', '.join(self.json_data.get('reference_documents', [])) or 'None'}",
            f"Additional Information: {self.json_data.get('additional_information', '').strip() or 'None'}"
        ]
        non_llm_summary = "\n\n".join(filter(None, summary_parts))

        core_record = {
            'process_id': process_id,
            'process_name': self.process_name,
            'doc_name': self.doc_name,
            'domain': self.json_data.get('domain', ''),
            'sub_domain': self.json_data.get('subdomain', ''), # Changed key
            'functional_area': self.json_data.get('functional_area', ''), # Assume these might exist
            'functional_subarea': self.json_data.get('functional_subarea', ''),
            'process_group': self.json_data.get('process_group', ''),
            'process_subgroup': self.json_data.get('process_subgroup', ''),
            'reference_documents': ', '.join(self.json_data.get('reference_documents', [])),
            'related_products': ', '.join(self.json_data.get('related_products', [])),
            'additional_information': self.json_data.get('additional_information', ''),
            'non_llm_summary': non_llm_summary
        }

        logger.info("      ✅ Core Record prepared successfully.")
        return core_record

    def _prepare_detailed_records(self, process_id: str) -> List[Dict[str, Any]]:
        """
        Prepares records for the detailed Azure Search index (steps).
        """
        print("   📑 Preparing Detailed Records (Steps)...")
        detailed_records = []

        # Step 0: Introduction/Overview
        logger.debug("      Adding Step 0 (Introduction/Overview)...")
        intro_content_parts = [
             f"Introduction: {self.json_data.get('introduction', '').strip()}",
             f"Short Description: {self.json_data.get('short_description', '').strip()}",
             f"Related Products: {', '.join(self.json_data.get('related_products', [])) or 'None'}",
             f"Reference Documents: {', '.join(self.json_data.get('reference_documents', [])) or 'None'}"
        ]
        intro_content = "\n\n".join(filter(None, intro_content_parts))

        intro_record = {
            'id': self.generate_step_id(0, "Introduction / Overview", intro_content),
            'process_id': process_id,
            'step_number': 0,
            'step_name': "Introduction / Overview",
            'step_content': intro_content,
            'documents_used': "", # Assume not applicable
            'systems_used': ""  # Assume not applicable
        }
        detailed_records.append(intro_record)
        logger.debug("         ✔️ Step 0 record created.")

        # Regular Steps
        steps_data = self.json_data.get('steps', [])
        logger.debug(f"      Processing {len(steps_data)} regular steps...")
        for i, step in enumerate(steps_data):
            step_num_str = step.get('step_number', str(i + 1))
            try:
                 step_num = int(step_num_str)
            except (ValueError, TypeError):
                 logger.warning(f"         ⚠️ Invalid step number '{step_num_str}' for step {i+1}. Using sequential {i+1}.")
                 step_num = i + 1

            step_name = step.get('step_name', f'Unnamed Step {step_num}')
            # Use 'step_description' as per format.json
            step_content = step.get('step_description', '').strip()

            record = {
                'id': self.generate_step_id(step_num, step_name, step_content),
                'process_id': process_id,
                'step_number': step_num,
                'step_name': step_name,
                'step_content': step_content,
                'documents_used': ', '.join(step.get('documents_used', [])),
                'systems_used': ', '.join(step.get('systems_used', []))
            }
            detailed_records.append(record)
            logger.debug(f"         ✔️ Step {step_num} record created ('{step_name[:30]}...')")

        logger.info(f"      ✅ Detailed Records prepared successfully ({len(detailed_records)} total records).")
        return detailed_records

    def prepare_for_upload(self) -> tuple[Dict[str, Any], List[Dict[str, Any]]]:
        """
        Prepares all records (core and detailed) from the JSON data for upload.
        """
        print("🚀 Preparing all records for upload...")
        process_id = self.generate_process_id()
        core_record = self._prepare_core_record(process_id)
        detailed_records = self._prepare_detailed_records(process_id)
        print("🏁 Upload preparation completed.")
        return core_record, detailed_records

#############################################################
# INGESTION CLASS (from ingestion.py)
#############################################################

class MultiProcessHandler:
    def __init__(self, json_paths: List[str], client_core: SearchClient, client_detail: SearchClient, oai_client: AzureOpenAI, embedding_model: str):
        """
        Initializes the handler for processing multiple JSON files and ingesting into Azure Search.
        """
        print(f"🚀 Initializing MultiProcessHandler for {len(json_paths)} JSON file(s)...")
        self.json_paths = json_paths
        self.client_core = client_core
        self.client_detail = client_detail
        self.oai_client = oai_client
        self.embedding_model = embedding_model # Get model from config
        logger.info(f"   🧠 Using embedding model: {self.embedding_model}")

    def process_all_documents(self) -> List[Dict[str, Any]]:
        """
        Processes multiple JSON documents, preparing core and detailed records for each.
        """
        print(f"🔄 Processing {len(self.json_paths)} JSON documents...")
        all_processed_records = []
        failed_files = []

        for json_path in tqdm(self.json_paths, desc="Processing JSON files"):
            print(f"\n   Processing file: {json_path}")
            try:
                # Use ProcessHandler (defined above) to structure data
                document_processor = ProcessHandler(json_path)
                core_record, detailed_records = document_processor.prepare_for_upload()

                all_processed_records.append({
                    'source_json': json_path, # Keep track of the source
                    'core': core_record,
                    'detailed': detailed_records
                })
                logger.info(f"      ✔️ Successfully prepared records from {os.path.basename(json_path)}")

            except FileNotFoundError:
                logger.error(f"      ❌ Error: File not found - {json_path}")
                failed_files.append(json_path)
            except json.JSONDecodeError as e:
                 logger.error(f"      ❌ Error: Invalid JSON in file - {json_path}: {e}")
                 failed_files.append(json_path)
            except Exception as e:
                logger.error(f"      ❌ Error processing {json_path}: {e}")
                logger.exception(f"Processing failed for {json_path}")
                failed_files.append(json_path)

        if failed_files:
             logger.warning(f"\n⚠️ Warning: Failed to process {len(failed_files)} JSON file(s):")
             for failed in failed_files: logger.warning(f"   - {failed}")

        print(f"🏁 Document processing phase complete. Prepared records for {len(all_processed_records)} files.")
        return all_processed_records

    def _generate_embeddings_batch(self, texts: List[str]) -> List[List[float]]:
        """Generates embeddings for a batch of texts, handling potential errors."""
        # Simplified batching - consider actual batch limits for the API
        embeddings = []
        valid_texts_indices = [i for i, text in enumerate(texts) if text and isinstance(text, str)]
        valid_texts = [texts[i] for i in valid_texts_indices]
        results = [[] for _ in range(len(texts))] # Initialize with empty lists

        if not valid_texts:
            return results

        try:
            logger.debug(f"      🧠 Generating embeddings for {len(valid_texts)} text(s)...")
            response = self.oai_client.embeddings.create(
                input=valid_texts,
                model=self.embedding_model
            )
            for i, idx in enumerate(valid_texts_indices):
                if response.data and i < len(response.data):
                     results[idx] = response.data[i].embedding
                else:
                     logger.warning(f"      ⚠️ Missing embedding data for text index {i} (original index {idx}).")
            logger.debug(f"      ✔️ Embeddings generated.")
        except Exception as e:
            logger.error(f"      ❌ Error generating embeddings batch: {e}")

        return results

    def upload_to_azure_search(self, all_processed_records: List[Dict[str, Any]]) -> None:
        """
        Uploads the processed and embedding-enriched records to Azure Search indexes.
        """
        if not all_processed_records:
            print(" M No records to upload.")
            return

        print(f"☁️ Starting upload process for {len(all_processed_records)} document(s) to Azure Search...")
        total_core_uploaded, total_detailed_uploaded = 0, 0
        total_core_failed, total_detailed_failed = 0, 0

        for record_set in tqdm(all_processed_records, desc="Uploading Records"):
            source_file = os.path.basename(record_set.get('source_json', 'Unknown Source'))
            print(f"\n   Uploading records for: {source_file}")

            core_record = record_set.get('core')
            detailed_records = record_set.get('detailed', [])

            # Enrich Core Record
            if core_record and 'non_llm_summary' in core_record:
                 summary_text = core_record['non_llm_summary']
                 embeddings = self._generate_embeddings_batch([summary_text])
                 core_record['embedding_summary'] = embeddings[0] if embeddings and embeddings[0] else None

            # Enrich Detailed Records
            if detailed_records:
                 step_names = [step.get('step_name', '') for step in detailed_records]
                 step_contents = [step.get('step_content', '') for step in detailed_records]
                 name_embeddings = self._generate_embeddings_batch(step_names)
                 content_embeddings = self._generate_embeddings_batch(step_contents)

                 for i, step in enumerate(detailed_records):
                      step['embedding_title'] = name_embeddings[i] if name_embeddings[i] else None
                      step['embedding_content'] = content_embeddings[i] if content_embeddings[i] else None
                      step['id'] = str(step.get('id', '')) # Ensure string IDs
                      step['process_id'] = str(step.get('process_id', ''))

            # Upload Core Record
            if core_record:
                 core_record['process_id'] = str(core_record.get('process_id', ''))
                 if core_record.get('embedding_summary') is None: core_record.pop('embedding_summary', None)
                 try:
                      logger.debug(f"      Uploading core record (ID: {core_record['process_id']})...")
                      upload_result = self.client_core.upload_documents(documents=[core_record])
                      if all(r.succeeded for r in upload_result):
                          logger.debug("         ✔️ Core record uploaded successfully.")
                          total_core_uploaded += 1
                      else:
                           total_core_failed += 1
                           logger.error(f"         ❌ Failed to upload core record for {source_file}.")
                           for r in upload_result:
                               if not r.succeeded: logger.error(f"            Error for key {r.key}: {r.error_message}")
                 except Exception as e:
                      total_core_failed += 1
                      logger.error(f"      ❌ Exception during core record upload for {source_file}: {e}")
                      logger.exception("Core upload failed.")

            # Upload Detailed Records
            if detailed_records:
                 cleaned_detailed_records = []
                 for step in detailed_records:
                     if step.get('embedding_title') is None: step.pop('embedding_title', None)
                     if step.get('embedding_content') is None: step.pop('embedding_content', None)
                     cleaned_detailed_records.append(step)
                 try:
                      logger.debug(f"      Uploading {len(cleaned_detailed_records)} detailed records...")
                      upload_result = self.client_detail.upload_documents(documents=cleaned_detailed_records)
                      success_count = sum(1 for r in upload_result if r.succeeded)
                      fail_count = len(cleaned_detailed_records) - success_count
                      total_detailed_uploaded += success_count
                      total_detailed_failed += fail_count
                      if fail_count == 0:
                          logger.debug(f"         ✔️ All {success_count} detailed records uploaded successfully.")
                      else:
                          logger.warning(f"         ⚠️ Uploaded {success_count} detailed records successfully, {fail_count} failed.")
                          for r in upload_result:
                              if not r.succeeded: logger.error(f"            Error for key {r.key}: {r.error_message}")
                 except Exception as e:
                      total_detailed_failed += len(detailed_records)
                      logger.error(f"      ❌ Exception during detailed records upload for {source_file}: {e}")
                      logger.exception("Detailed upload failed.")

        print("\n🏁 Upload process finished.")
        print(f"   📊 Summary: Core Records (Uploaded: {total_core_uploaded}, Failed: {total_core_failed})")
        print(f"   📊 Summary: Detailed Records (Uploaded: {total_detailed_uploaded}, Failed: {total_detailed_failed})")

#############################################################
# AZURE SEARCH INDEX CREATION FUNCTIONS (from index_creation.py)
#############################################################

def create_vector_search_configuration() -> VectorSearch:
    """Create vector search configuration."""
    logger.info("Creating vector search configuration...")
    # Using ExhaustiveKnn based on index_creation.py example
    return VectorSearch(
        algorithms=[
            ExhaustiveKnnAlgorithmConfiguration(
                name='vector-config',
                kind='exhaustiveKnn',
                parameters=ExhaustiveKnnParameters(metric="cosine")
            )
        ],
        profiles=[
            VectorSearchProfile(
                name='vector-search-profile',
                algorithm_configuration_name='vector-config',
            )
        ]
    )

def create_enhanced_core_df_index(index_name: str, vector_config: VectorSearch) -> SearchIndex:
    """Create index schema for core document dataframe."""
    logger.info(f"Defining schema for core index: {index_name}")
    fields=[
        # Key and Core Identifiers
        SimpleField(name="process_id", type=SearchFieldDataType.String, key=True, filterable=True, sortable=True, searchable=True, retrievable=True), # searchable=True added based on combined_pipeline
        SearchableField(name="process_name", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, sortable=True, analyzer_name='el.lucene'), # sortable=True added
        SearchableField(name="doc_name", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, analyzer_name='el.lucene'),

        # Hierarchical/Categorical Fields
        SearchableField(name="domain", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, facetable=True, analyzer_name='el.lucene'), # facetable=True added
        SearchableField(name="sub_domain", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, facetable=True, analyzer_name='el.lucene'), # facetable=True added
        SearchableField(name="functional_area", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, facetable=True, analyzer_name='el.lucene'), # Added filterable/facetable
        SearchableField(name="functional_subarea", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, facetable=True, analyzer_name='el.lucene'), # Added filterable/facetable
        SearchableField(name="process_group", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, facetable=True, analyzer_name='el.lucene'), # Added filterable/facetable
        SearchableField(name="process_subgroup", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, facetable=True, analyzer_name='el.lucene'), # Added filterable/facetable

        # Content and Metadata Fields (using String for lists/arrays)
        SearchableField(name="reference_documents", type=SearchFieldDataType.String, searchable=True, retrievable=True, analyzer_name='el.lucene'),
        SearchableField(name="related_products", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, analyzer_name='el.lucene'), # filterable=True added
        SearchableField(name="additional_information", type=SearchFieldDataType.String, searchable=True, retrievable=True, analyzer_name='el.lucene'),
        SearchableField(name="non_llm_summary", type=SearchFieldDataType.String, searchable=True, retrievable=True, analyzer_name='el.lucene'),

        # Vector Field
        SearchField(
            name="embedding_summary", # Field containing the vector embedding
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True, # Enables vector search
            vector_search_dimensions=3072, # Dimension (e.g., text-embedding-3-large)
            vector_search_profile_name="vector-search-profile" # Link profile
        ),
    ]
    semantic_config = SemanticConfiguration(
        name="enhanced-core-df-semantic-config",
        prioritized_fields=SemanticPrioritizedFields(
            title_field=SemanticField(field_name="process_name"),
            content_fields=[SemanticField(field_name="non_llm_summary")],
            keywords_fields=[
                SemanticField(field_name="domain"), SemanticField(field_name="sub_domain"),
                SemanticField(field_name="functional_area"), SemanticField(field_name="functional_subarea"),
                SemanticField(field_name="process_group"), SemanticField(field_name="process_subgroup"),
                SemanticField(field_name="related_products") # Added based on combined_pipeline
            ]
        )
    )
    return SearchIndex(
        name=index_name,
        fields=fields,
        semantic_search=SemanticSearch(configurations=[semantic_config]),
        vector_search=vector_config
    )

def create_enhanced_detailed_df_index(index_name: str, vector_config: VectorSearch) -> SearchIndex:
    """Create index schema for detailed document dataframe (steps)."""
    logger.info(f"Defining schema for detailed index: {index_name}")
    fields=[
        # Key and Linking Fields
        SimpleField(name="id", type=SearchFieldDataType.String, key=True, filterable=True, searchable=True, retrievable=True), # searchable=True added
        SimpleField(name="process_id", type=SearchFieldDataType.String, filterable=True, sortable=True, searchable=True, retrievable=True), # sortable=True, searchable=True added

        # Step Specific Fields
        SimpleField(name="step_number", type=SearchFieldDataType.Int64, filterable=True, sortable=True, retrievable=True),
        SearchableField(name="step_name", type=SearchFieldDataType.String, searchable=True, retrievable=True, analyzer_name='el.lucene'),
        SearchableField(name="step_content", type=SearchFieldDataType.String, searchable=True, retrievable=True, analyzer_name='el.lucene'),
        SearchableField(name="documents_used", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, analyzer_name='el.lucene'), # filterable=True added
        SearchableField(name="systems_used", type=SearchFieldDataType.String, searchable=True, retrievable=True, filterable=True, analyzer_name='el.lucene'), # filterable=True added

        # Vector Fields
        SearchField(
            name="embedding_title",
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=3072,
            vector_search_profile_name="vector-search-profile"
        ),
        SearchField(
            name="embedding_content",
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=3072,
            vector_search_profile_name="vector-search-profile"
        )
    ]
    semantic_config = SemanticConfiguration(
        name="enhanced-detailed-df-semantic-config",
        prioritized_fields=SemanticPrioritizedFields(
            title_field=SemanticField(field_name="step_name"),
            content_fields=[SemanticField(field_name="step_content")],
            keywords_fields=[ # Added based on combined_pipeline
                 SemanticField(field_name="documents_used"),
                 SemanticField(field_name="systems_used")
            ]
        )
    )
    return SearchIndex(
        name=index_name,
        fields=fields,
        semantic_search=SemanticSearch(configurations=[semantic_config]),
        vector_search=vector_config
    )

def list_indexes(index_client: SearchIndexClient):
    """List all indexes in the Azure AI Search service."""
    try:
        logger.info("Listing existing indexes...")
        indexes = index_client.list_indexes()
        index_names = [index.name for index in indexes]
        if index_names:
            logger.info(f"Found {len(index_names)} indexes:")
            for name in index_names: logger.info(f"- {name}")
        else:
            logger.info("No indexes found in the service.")
    except Exception as e:
        logger.error(f"An error occurred while listing indexes: {e}")


def manage_azure_search_indexes(index_client: SearchIndexClient, core_index_name: str, detailed_index_name: str, recreate: bool = False):
    """Manage Azure Cognitive Search indexes (Create or Recreate)."""
    try:
        logger.info("Starting Azure Search index management...")
        list_indexes(index_client) # List before operation

        vector_config = create_vector_search_configuration()
        core_df_index = create_enhanced_core_df_index(core_index_name, vector_config)
        detailed_df_index = create_enhanced_detailed_df_index(detailed_index_name, vector_config)
        indexes_to_create = {core_index_name: core_df_index, detailed_index_name: detailed_df_index}

        if recreate:
            logger.warning("Recreate flag is True. Attempting to delete existing indexes...")
            for index_name in indexes_to_create:
                try:
                    index_client.delete_index(index_name)
                    logger.info(f"Deleted index '{index_name}' successfully!")
                except Exception as e:
                    logger.warning(f"Could not delete index '{index_name}': {e} (Might not exist)")
            logger.info("Existing indexes after deletion attempt:")
            list_indexes(index_client)

        logger.info("Attempting to create or update indexes...")
        for index_name, index_definition in indexes_to_create.items():
            try:
                index_client.create_or_update_index(index_definition)
                logger.info(f"Index '{index_name}' created or updated successfully!")
            except Exception as e:
                logger.error(f"Error creating/updating index '{index_name}': {e}")

        logger.info("Index management process completed.")
        logger.info("Final list of indexes:")
        list_indexes(index_client)
    except Exception as e:
        logger.critical(f"A critical error occurred during index management: {e}")
        logger.exception("Index management failed.")


#############################################################
# MAIN EXECUTION LOGIC (Adapted from main.py)
#############################################################

def run_pipeline(config: Dict[str, Any], manage_indexes: bool = False, recreate_indexes: bool = False):
    """
    Main function to orchestrate the entire knowledge pipeline process.

    Args:
        config (dict): Loaded configuration dictionary.
        manage_indexes (bool): If True, run index management before processing.
        recreate_indexes (bool): If True and manage_indexes is True, delete existing indexes first.
    """
    global clients # Use the global clients variable

    print("🚀 Starting Knowledge Pipeline...")

    # --- Initialization ---
    clients = initialize_clients(config)
    if not clients:
        print("❌ Pipeline aborted due to client initialization failure.")
        return

    # Load the JSON format structure needed by DocParsing
    try:
        json_format_structure = load_format(config["FORMAT_JSON_PATH"])
    except Exception as e:
        print(f"❌ Error loading format JSON from {config['FORMAT_JSON_PATH']}: {e}. Aborting.")
        return

    # --- Optional: Index Management ---
    if manage_indexes:
        print("\n🛠️ Starting Index Management...")
        manage_azure_search_indexes(
            index_client=clients["search_index_client"],
            core_index_name=config["CORE_INDEX_NAME"],
            detailed_index_name=config["DETAILED_INDEX_NAME"],
            recreate=recreate_indexes
        )
        print("🛠️ Index Management Complete.")
        # Decide if you want to exit after index management or continue processing
        # Consider adding a command-line flag for this behaviour
        # return # Example: exit after managing indexes

    # --- File Discovery and Filtering ---
    print("\n🔍 Discovering files in Azure Blob Storage...")
    processed_files = load_processed_files()
    # failed_files_log = load_failed_files() # Consider if retry logic is needed

    files_to_process = []
    try:
        # Use the container client obtained during initialization
        blob_container_client = clients["container_client"]
        blob_list = blob_container_client.list_blobs() # List blobs

        for blob in blob_list:
            if blob.name.lower().endswith(".docx") and blob.name not in processed_files:
                 folder_path = os.path.dirname(blob.name)
                 file_name = os.path.basename(blob.name)
                 # Store blob path for direct access later
                 files_to_process.append({'folder': folder_path, 'name': file_name, 'blob_path': blob.name})
            elif blob.name in processed_files:
                 logger.info(f"   Skipping already processed file: {blob.name}")

        print(f"   Found {len(files_to_process)} new DOCX files to process.")

    except Exception as e:
        print(f"❌ Error listing blobs in container '{config['CONTAINER_NAME']}': {e}")
        logger.exception("Blob listing failed.")
        return # Abort if file discovery fails

    # --- Processing Loop ---
    if not files_to_process:
        print("\n✅ No new files to process. Pipeline finished.")
        return

    print(f"\n⚙️ Processing {len(files_to_process)} documents...")
    all_extracted_json_paths = [] # Collect paths of JSONs generated by DocParsing

    for file_info in tqdm(files_to_process, desc="Processing Documents"):
        folder = file_info['folder']
        file_name = file_info['name']
        blob_path = file_info['blob_path'] # Use the direct blob path
        doc_name_no_ext = os.path.splitext(file_name)[0]

        print(f"\n   Processing: {blob_path}")

        # Clean temp directory *before* processing each DOCX
        # This ensures isolation but check if intermediate JSONs need persistence across failures
        logger.info(f"      Cleaning temporary directory: {config['TEMP_PATH']}")
        delete_files_in_directory(config['TEMP_PATH'])

        try:
            # 1. Download DOCX content
            print(f"      Downloading {blob_path}...")
            blob_client = blob_container_client.get_blob_client(blob_path)
            blob_content = blob_client.download_blob().readall()
            byte_stream = BytesIO(blob_content)
            docx_document = Document(byte_stream)
            print("         ✔️ Downloaded and opened successfully.")

            # 2. Parse DOCX to JSON using DocParsing
            print("      Parsing document with DocParsing...")
            parser = DocParsing(
                doc_instance=docx_document,
                client=clients["azure_oai_client"],
                json_format=json_format_structure,
                domain="-", # Or derive domain from folder
                sub_domain=folder if folder else "root", # Use folder as sub-domain
                model_name=config["MODEL_NAME"],
                doc_name=doc_name_no_ext
            )
            # This method now handles extraction, AI call, and saving JSONs
            parser.doc_to_json(output_dir=config['TEMP_PATH'])

            # 3. Collect generated JSON file paths for this DOCX
            generated_jsons = get_all_files_in_directory(config['TEMP_PATH'])
            if generated_jsons:
                 logger.info(f"      ✔️ DocParsing generated {len(generated_jsons)} JSON file(s) in {config['TEMP_PATH']}.")
                 all_extracted_json_paths.extend(generated_jsons)
                 # Mark original DOCX blob_path as processed *only after successful parsing*
                 save_processed_file(blob_path) # Save the full blob path
                 logger.info(f"      ✔️ Marked '{blob_path}' as processed (parsing stage).")
            else:
                 logger.warning(f"      ⚠️ DocParsing did not generate any JSON files for {blob_path}. Skipping ingestion.")
                 log_failed_file(f"{blob_path} - DocParsing generated no JSONs")

        except Exception as e:
            print(f"❌ Error processing document {blob_path}: {e}")
            logger.exception(f"Failed processing {blob_path}")
            log_failed_file(f"{blob_path} - Error: {e}")
            continue # Continue to the next file

    # --- Ingestion Phase ---
    if not all_extracted_json_paths:
         print("\nℹ️ No JSON files were generated by DocParsing. Nothing to ingest.")
    else:
         print(f"\n🚚 Starting ingestion phase for {len(all_extracted_json_paths)} generated JSON file(s)...")
         try:
              ingestion_handler = MultiProcessHandler(
                   json_paths=all_extracted_json_paths,
                   client_core=clients["search_client_core"],
                   client_detail=clients["search_client_detail"],
                   oai_client=clients["azure_oai_client"],
                   embedding_model=config["MODEL_NAME"]
              )
              prepared_records = ingestion_handler.process_all_documents()
              if prepared_records:
                   ingestion_handler.upload_to_azure_search(prepared_records)
              else:
                   print("   ⚠️ No records were successfully prepared. Skipping Azure Search upload.")
         except Exception as e:
              print(f"❌ An error occurred during the ingestion phase: {e}")
              logger.exception("Ingestion phase failed.")

    print("\n🏁 Knowledge Pipeline finished.")


#############################################################
# COMMAND LINE INTERFACE (CLI) SETUP
#############################################################

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Knowledge Pipeline for processing DOCX files from Azure Blob Storage.")

    parser.add_argument(
        '--manage-indexes',
        action='store_true',
        help='Run index management tasks (list, create, or recreate indexes) before processing files.'
    )
    parser.add_argument(
        '--recreate-indexes',
        action='store_true',
        help='If --manage-indexes is set, delete existing indexes before creating new ones. Use with caution!'
    )
    # Add more arguments if needed, e.g., to specify config file path, temp dir, etc.
    # parser.add_argument('--config', type=str, default='.env', help='Path to the environment configuration file.')
    # parser.add_argument('--temp-dir', type=str, help='Override temporary directory path.')

    args = parser.parse_args()

    # Load configuration first
    pipeline_config = load_configuration()

    # Override config with CLI args if provided (example for temp-dir)
    # if args.temp_dir:
    #     pipeline_config["TEMP_PATH"] = args.temp_dir
    #     logger.info(f"Overriding TEMP_PATH with CLI argument: {args.temp_dir}")

    # Run the main pipeline function with CLI arguments
    run_pipeline(
        config=pipeline_config,
        manage_indexes=args.manage_indexes,
        recreate_indexes=args.recreate_indexes
    )