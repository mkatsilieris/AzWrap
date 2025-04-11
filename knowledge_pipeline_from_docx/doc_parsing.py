import os
import json
import re
from openai import AzureOpenAI
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from tqdm import tqdm

class DocParsing:
    def __init__(self, doc_instance, client, json_file, domain, sub_domain, model_name, doc_name):
        """
        Initialize the DocParsing class with necessary parameters.
        
        Parameters:
            doc_instance: Document object to be parsed
            client: Azure OpenAI client for AI processing
            json_file: Template for the JSON structure
            domain: Domain category for the document
            sub_domain: Sub-domain category for the document
            model_name: Name of the AI model to use
            doc_name: Name of the document being processed
        """
        print("🚀 Initializing DocParsing class...")
        self.client = client 
        self.doc = doc_instance  
        self.doc_name = None  
        self.format = json_file
        self.domain = domain
        self.model_name = model_name
        self.sub_domain = sub_domain
        self.doc_name = doc_name

        
    def get_section_header_lines(self, section):
        """
        Extract all text lines from a section's header.
        
        Collects text from paragraphs and tables in the header section.
        
        Parameters:
            section: The document section to extract header text from
            
        Returns:
            List of text lines from the header
        """
        try:
            if not section or not section.header:
                return []

            lines = []
            # Gather paragraph text from the header
            for paragraph in section.header.paragraphs:
                txt = paragraph.text.strip()
                if txt:
                    lines.append(txt)

            # Gather table cell text from the header (if any)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_txt = cell.text.strip()
                        if cell_txt:
                            lines.append(cell_txt)
            return lines
        except Exception as e:
            return []

    def parse_header_lines(self, header_lines):
        """
        Parse header lines to extract the process title.
        
        Analyzes header text to identify process numbers and titles,
        filtering out metadata like edition info and page numbers.
        
        Parameters:
            header_lines: List of text lines from a header
            
        Returns:
            String containing the process title or "Metadata" if none found
        """
        # Skip empty lists
        if not header_lines:
            return "Metadata"
            
        # Pattern for process numbers 
        number_pattern = re.compile(r'^\d+(?:\.\d+)*\.$')
        
        # Look for process number and title pattern in each line
        for line in header_lines:
            line_stripped = line.strip()
            
            # Skip metadata lines (edition info, page numbers)
            if line_stripped.startswith("Εκδ.") or "Σελ." in line_stripped:
                continue
                
            # Check if line contains a process number with title
            if "\t" in line_stripped:
                # Format: "1.1.1.1.\tTitle"
                parts = line_stripped.split("\t", 1)
                if number_pattern.match(parts[0].strip()) and len(parts) > 1:
                    return parts[1].strip()
                    
            # Check if line is just a process number
            elif number_pattern.match(line_stripped):
                # Look for title in the next lines
                for next_line in header_lines[header_lines.index(line) + 1:]:
                    next_stripped = next_line.strip()
                    if next_stripped and not next_stripped.startswith("Εκδ.") and "Σελ." not in next_stripped:
                        return next_stripped
    
        # If no match found, return "Metadata"
        return "Metadata"

    def extract_header_info(self, section):
        """
        Extract process title from section header.
        
        Combines header line extraction and parsing to get a process title.
        
        Parameters:
            section: The document section to extract title from
            
        Returns:
            String containing the process title or None if extraction fails
        """
        try:
            if not section or not section.header:
                return None

            lines = self.get_section_header_lines(section)
            header_title = self.parse_header_lines(lines)
            
            # For debugging
            if header_title == "Metadata" and lines:
                print(f"")
                
            return header_title
        except Exception as e:
            print(f"Error extracting header info: {e}")
            return None


    def iterate_block_items_with_section(self, doc):
        """
        Iterate through document blocks (paragraphs and tables) with their section indices.
        
        Tracks section changes while iterating through elements.
        
        Parameters:
            doc: The document to iterate through
            
        Yields:
            Tuple of (section_index, block) where block is either a Paragraph or Table
        """
        parent_elm = doc._element.body
        current_section_index = 0

        for child in parent_elm.iterchildren():
            if child.tag.endswith("p"):
                paragraph = Paragraph(child, doc)
                yield current_section_index, paragraph
                # Check for a new section
                if child.xpath(".//w:sectPr"):
                    current_section_index += 1

            elif child.tag.endswith("tbl"):
                table = Table(child, doc)
                yield current_section_index, table

    def extract_table_data(self, table):
        """
        Extract text data from a table.
        
        Converts table rows into strings with cells joined by ' - '.
        
        Parameters:
            table: Table object to extract data from
            
        Returns:
            String containing all table data with rows separated by newlines
        """
        data = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            # row as string
            if row_cells:
                row_string = ' - '.join(row_cells)  # Join cells with a space or another separator
                data.append(row_string)
        return '\n'.join(data)

    
    def is_single_process(self, doc, doc_name):
        """
        Check if document is a single process document.
        
        Determines if all sections have the same process title or no process titles.
        
        Parameters:
            doc: Document object to check
            doc_name: Name of the document
            
        Returns:
            Tuple of (is_single_process: bool, title: str or None)
        """
        print("🔎 Checking if single process document...")
        
        # Get all section headers
        section_headers = set()
        non_metadata_sections = 0
        
        for section_index in range(len(doc.sections)):
            section = doc.sections[section_index]
            header_title = self.extract_header_info(section)
            
            if header_title and header_title != "Metadata":
                section_headers.add(header_title)
                non_metadata_sections += 1
        
        # A document is single-process if only one unique header title exists
        # or if there are zero unique titles (simple document)
        if len(section_headers) <= 1:
            # Single title or no titles at all
            title = list(section_headers)[0] if section_headers else doc_name
            return True, title
        else:
            # Multiple different titles - multi-process document
            return False, None
    
    def extract_data(self):
        """
        Extract content from the document based on section headers.
        
        Processes document differently depending on whether it's a single
        or multi-process document.
        
        Returns:
            Dictionary with keys as formatted headers and values as section content
        """
        doc_name = self.doc_name
        doc = self.doc
        data_dict = {}
        
        # Check if single or multiple process document
        is_single_process_bool, single_header = self.is_single_process(doc, doc_name)
        
        if is_single_process_bool:
            # Single process document processing
            header_title = f"{single_header}_single_process"
            data_dict[header_title] = []
            for section_index, block in self.iterate_block_items_with_section(doc):
                if isinstance(block, Paragraph) and block.text and block.text.strip():
                    data_dict[header_title].append(block.text)
        else:
            # Multi-process document processing
            # First identify all unique sections
            for section_index in range(len(doc.sections)):
                header_title = self.extract_header_info(doc.sections[section_index])
                if not header_title:
                    print(f"Warning: Section {section_index} has no extractable header")
                    header_title = "Unknown"
                    
                formated_header = f"{doc_name}_header_{header_title}"
                if formated_header not in data_dict:
                    data_dict[formated_header] = []
            
            # Then collect content for each section
            for section_index, block in self.iterate_block_items_with_section(doc):
                if section_index < len(doc.sections):
                    header_title = self.extract_header_info(doc.sections[section_index])
                    if not header_title:
                        header_title = "Unknown"
                        
                    formated_header = f"{doc_name}_header_{header_title}"
                    if isinstance(block, Paragraph) and block.text and block.text.strip():
                        if formated_header in data_dict:
                            data_dict[formated_header].append(block.text)
                else:
                    print(f"Warning: Block referenced section_index {section_index} which exceeds section count {len(doc.sections)}")
        
        # Join the content for each section
        for key, value in data_dict.items():
            data_dict[key] = "\n".join(value)
        
        return data_dict

    def update_json(self, data_format, content, name):
        """
        Use AI to parse document content into structured JSON format.
        
        Sends document content to the AI model with specific prompts
        to extract and organize information.
        
        Parameters:
            data_format: JSON structure template
            content: Document content to parse
            name: Document name for context
            
        Returns:
            JSON string containing the parsed content
        """
        print("🤖 Updating JSON using AI...")
        prompt = (
            "Parse the provided information about a specific process from the document and fill in the JSON structure below. "
            "Do not summarize, omit, or modify any details. Simply extract and organize the provided data into the corresponding fields of the JSON. "
            "There are more than one step and you have to include all of them.The step description has to be the whole text till the next step name"
            "Ensure every relevant detail is included without altering the content. "
            "The JSON format should follow this structure and include all fields, even if some of them are not present in the content (leave them empty or null if necessary):\n"
            "To make it clear the content you generate will be ONLY THE CONTENT of a json no \n nothing.The first character { and the last character should be }"
            "Your response should be ONLY a JSON file content ready to be stored as json without other processing, with the exact format as shown above."
        )
        output_llm = self.client.chat.completions.create(
        model= self.model_name,
        messages=[
                {"role": "system", "content": prompt + "\n" + f"{data_format}"} , 
                {"role": "user", "content": f"Document name : {name} \n Content to parse: {content}"}
            ],
        temperature = 0
        )
        return (output_llm.choices[0].message.content)

    def process_and_generate_json(self, response_str, output_file):
        """
        Process AI response and save it as a properly formatted JSON file.
        
        Cleans the AI response, adds domain and subdomain information,
        and writes the result to a file.
        
        Parameters:
            response_str: AI response string containing JSON data
            output_file: Path to save the output JSON file
        """
        print("💾 Processing and generating JSON file...")
        # Remove ```json from start and ``` from end if present
        cleaned_response = re.sub(r"^```json\s*|\s*```$", "", response_str.strip())

        try:
            # Parse the cleaned response into a Python dictionary
            json_data = json.loads(cleaned_response)

            # Ensure process_name exists and insert domain & subdomain right after
            if "process_name" in json_data:
                ordered_data = {
                    "doc_name": self.doc_name,
                    "process_name": json_data["process_name"],
                    "domain": self.domain,
                    "subdomain": self.sub_domain,
                    **{k: v for k, v in json_data.items() if k != "process_name"}
                }
            else:
                # If process_name is missing, just add domain and subdomain normally
                ordered_data = {"domain": self.domain, "subdomain": self.sub_domain, **json_data}

            # Write updated JSON to the output file
            with open(output_file, 'w', encoding='utf-8') as file:
                json.dump(ordered_data, file, indent=4, ensure_ascii=False)

            print(f"✅ JSON data successfully written to {output_file}")

        except json.JSONDecodeError as e:
            print(f"❌ Error decoding JSON: {e}")
        except Exception as e:
            print(f"❌ Unexpected error: {e}")

    def doc_to_json(self, doc_name=None, output_dir="temp_json"):
        """
        Main method to convert a document to JSON files.
        
        Extracts data from the document, processes each section,
        and saves the results as JSON files.
        
        Parameters:
            doc_name: Optional document name override
            output_dir: Directory to save output JSON files
            
        Returns:
            None
        """
        print("🚢 Starting document to JSON conversion...")
        self.doc_name = self.doc_name
        data_dict = self.extract_data()
        
        # Ensure the json directory exists
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        for i in data_dict:
            if "Metadata" not in i:
                name = i.replace("/" , "_")
                path = f"{output_dir}/{name}"
                print(path)
                content = self.update_json(self.format, data_dict[i], self.doc_name)
                self.process_and_generate_json(content, f"{path}.json")
        print("🏁 Document to JSON conversion completed!")
        return

# Example usage:

# doc = Document(r"C:\Users\agkithko\OneDrive - Netcompany\Desktop\ΚΑΤΑΝΑΛΩΤΙΚΑ_ΔΙΑΔΙΚΑΣΙΕΣ-20250408T073725Z-001\ΚΑΤΑΝΑΛΩΤΙΚΑ_ΔΙΑΔΙΚΑΣΙΕΣ\J66_001_2024-ΔΙΕΡΕΥΝΗΣΗ ΑΝΑΓΚΩΝ ΠΕΛΑΤΗ ΕΞΟΙΚΟΝΟΜΩ_WF.docx")
# parser = DocParsing(doc_instance=doc, format_path=r"C:\Users\agkithko\OneDrive - Netcompany\Desktop\AzWrap-1\temp_json", domain="your_domain", sub_domain="your_sub_domain")
# parser.doc_to_json(doc_name="optional_custom_name")